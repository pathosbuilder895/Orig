"""
api.py — FastAPI application for the dashboard demo and pilot-compatible server.

Endpoints
─────────
GET  /health
GET  /students                                      list all student IDs
GET  /students/{id}                                 student state summary
POST /students/{id}/baseline                        add a baseline sample (text)
POST /students/{id}/baseline/upload-batch           add multiple files as baseline
POST /students/{id}/score                           score a submission → Layer 7
POST /students/{id}/upload                          extract text from a single file
POST /import/courses/{course_id}/turnitin-csv       import Turnitin CSV export
POST /canvas/baseline/{id}/list-canvas-submissions  list past Canvas submissions for student
POST /canvas/baseline/{id}/import-baseline          import selected Canvas submissions as baseline

In demo mode, anonymous sandbox access remains available for the seeded sales
demo. In pilot/production modes, real tenant data is protected by the Principal
tenant-isolation middleware, stable SECRET_KEY requirement, locked CORS, and
guarded destructive operations.
"""

from __future__ import annotations

import csv
import hmac
import io
import json
import logging
import os
import urllib.parse
import uuid
from contextlib import asynccontextmanager

from fastapi import FastAPI, Form, HTTPException, Request, UploadFile, File
from typing import List, Optional
from fastapi.middleware.cors import CORSMiddleware

from .schemas import (
    AddSampleRequest,
    ScoreSubmissionRequest,
    Layer7OutputResponse,
    StudentStateResponse,
    SampleSummary,
    HealthResponse,
    AuthorshipSignalOut,
    TrajectoryConformanceOut,
    FeatureContributionOut,
    EntanglementAnomalyOut,
    InterferenceDecompositionOut,
    BaselineConfidenceOut,
    DomainSignalOut,
    RecommendedActionOut,
    TensionArcOut,
    ContextManifestOut,
    ScoringReportOut,
    BlendDetectionRequest,
    BlendResultOut,
    WindowScoreOut,
    DriftResultOut,
    DriftPendingResponse,
    DriftRebaselineResponse,
    ManifestListItem,
    ManifestListResponse,
    ManifestStatsResponse,
    CorrectionRequest,
    CorrectionResponse,
    CorrectionListResponse,
    TestScoreRequest,
    TestScoreResponse,
    DatasetInfo,
    CalibrationRunRequest,
    CalibrationRunSummary,
    CalibrationRunDetail,
    CalibrationRunListResponse,
    CalibrationRunCreatedResponse,
    SuggestionItem,
    SuggestionsResponse,
    ApplyThresholdsRequest,
    TunedThresholdsRecord,
    TunedThresholdsListResponse,
)
from .tension_arc import analyze_tension_arc, update_student_baseline_kappa
from .features.pipeline import extract_features, feature_vector
from .quantum.state import BaselineSample
from .quantum.scoring import score as quantum_score
from .constants import AUTH_WEIGHTS, FEATURE_DIM
from . import store
from . import baseline_requests
from . import bbook_client
from . import student_auth
from . import principal as principal_mod
from . import users as users_mod
from .repository import get_repository
from fastapi.responses import JSONResponse, RedirectResponse, HTMLResponse

# NOTE: .env is loaded by the run.py entrypoint (not at import) so importing the
# app in tests/other contexts never pollutes os.environ for the v1 Settings.

# Deployment mode for the legacy/demo app: "demo" (default, zero-login sandbox),
# "pilot", or "production". Controls CORS defaults, the SECRET_KEY fail-fast,
# and security headers. Distinct from the v1 app's ENVIRONMENT setting.
ORIGINAL_ENV = os.environ.get("ORIGINAL_ENV", "demo").strip().lower()
_IS_REAL_DEPLOY = ORIGINAL_ENV in ("pilot", "staging", "production")

@asynccontextmanager
async def lifespan(app: FastAPI):
    _log = logging.getLogger(__name__)
    if not _secret_key_pinned:
        if _IS_REAL_DEPLOY:
            # Fail closed: a random per-process key silently invalidates every
            # token on restart and is unacceptable outside the demo sandbox.
            raise RuntimeError(
                f"ORIGINAL_ENV={ORIGINAL_ENV} requires a stable SECRET_KEY. "
                "Set it in the environment or .env: "
                "python -c \"import secrets; print(secrets.token_urlsafe(64))\""
            )
        _log.warning(
            "SECRET_KEY is not set — using a per-process random value. "
            "JWTs will be invalidated on every restart. "
            "Set SECRET_KEY in your environment or .env file for a stable key: "
            "  python -c \"import secrets; print(secrets.token_urlsafe(64))\""
        )
    else:
        _log.info("SECRET_KEY is pinned from environment — JWT tokens survive restarts.")
    _log.info("ORIGINAL_ENV=%s — CORS=%s", ORIGINAL_ENV, _ALLOWED_ORIGINS)
    _log.info(
        "GUARD_DESTRUCTIVE=%s — destructive endpoints are %s.",
        _GUARD_DESTRUCTIVE,
        "GUARDED (X-Guard-Token required)" if _GUARD_DESTRUCTIVE else "open (demo mode)",
    )
    yield


app = FastAPI(
    title="Original — Authorship Integrity API",
    version="0.1.0",
    description="Quantum stylometric authorship analysis for seminary submissions.",
    lifespan=lifespan,
)

# CORS: demo allows any origin; pilot/production must list origins explicitly
# via ALLOWED_ORIGINS (comma-separated). Falls back to "*" only in demo.
def _resolve_allowed_origins():
    raw = os.environ.get("ALLOWED_ORIGINS", "").strip()
    if raw:
        return [o.strip() for o in raw.split(",") if o.strip()]
    if _IS_REAL_DEPLOY:
        return []  # locked down: no origin allowed until configured
    return ["*"]


_ALLOWED_ORIGINS = _resolve_allowed_origins()

app.add_middleware(
    CORSMiddleware,
    allow_origins=_ALLOWED_ORIGINS,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Security headers ──────────────────────────────────────────────────────────
# Always-safe headers on every response. HSTS is opt-in (ENABLE_HSTS=1) since it
# only makes sense once TLS terminates in front of the app. X-Frame-Options is
# SAMEORIGIN (not DENY) so LTI launches can render inside an LMS that we allow
# via a CSP frame-ancestors directive at deploy time.
_ENABLE_HSTS = os.environ.get("ENABLE_HSTS", "0") == "1"


@app.middleware("http")
async def security_headers(request: "Request", call_next):
    resp = await call_next(request)
    resp.headers.setdefault("X-Content-Type-Options", "nosniff")
    resp.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    # LTI launches render inside the LMS iframe, so don't frame-block them.
    # (Restrict embedders at deploy time via a CSP frame-ancestors directive.)
    if not request.url.path.startswith("/lti/"):
        resp.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
    if _ENABLE_HSTS:
        resp.headers.setdefault(
            "Strict-Transport-Security", "max-age=31536000; includeSubDomains"
        )
    return resp


# ── Tenant-isolation middleware (ADR-003, Phase 1) ────────────────────────────
# Resolves the request principal once and enforces tenant boundaries on every
# student-scoped path in ONE place. Additive by construction: with no
# credentials the principal is the anonymous demo principal, which is allowed
# over flat ids + demo-environment tenants — i.e. today's demo is unchanged.
# Real (pilot/production) tenant data is only reachable by an authenticated
# principal of that tenant (or a super/operator role). See original/principal.py.
@app.middleware("http")
async def tenant_isolation(request: "Request", call_next):
    principal = principal_mod.resolve_principal(request)
    request.state.principal = principal
    scoped_id = principal_mod.extract_scoped_id(request.url.path)
    if scoped_id is not None:
        try:
            principal_mod.assert_student_access(principal, scoped_id)
        except principal_mod.TenantAccessError:
            return JSONResponse(
                status_code=403,
                content={"detail": "Cross-tenant access denied."},
            )
    return await call_next(request)

# ── Startup: SECRET_KEY stability check ───────────────────────────────────────
# Warn operators if SECRET_KEY is not pinned in the environment.
# A random per-process key means all JWTs issued by a prior process are
# immediately invalidated on restart — silent auth breakage that's hard to
# debug in a real deployment. Demo mode is expected to be ephemeral; any
# other environment should pin a stable key via the SECRET_KEY env var.

_secret_key_pinned: bool = bool(os.environ.get("SECRET_KEY", ""))

# ── GUARD_DESTRUCTIVE flag ────────────────────────────────────────────────────
# When GUARD_DESTRUCTIVE=1, three high-risk endpoints (student deletion,
# calibration threshold apply, and baseline-request list) require a
# matching X-Guard-Token header.  This lets pilot-mode deployments protect
# dangerous operations without a full JWT/RBAC stack.
#
# Demo mode leaves this unset so the frontend works without credentials.
# Pilot/production operators should:
#   1. Set MAINTENANCE_TOKEN to a strong random string
#   2. Set GUARD_DESTRUCTIVE=1
#   3. The X-Guard-Token header value must equal MAINTENANCE_TOKEN

_GUARD_DESTRUCTIVE: bool = os.environ.get("GUARD_DESTRUCTIVE", "0") == "1"


def _repo():
    """The persistence Repository for this environment (ADR-002 seam)."""
    return get_repository(os.environ.get("ENVIRONMENT", "demo"))


def _require_guard(request: "Request") -> None:
    """
    Raise 403 if GUARD_DESTRUCTIVE is on and the request lacks the correct
    X-Guard-Token header.  Call at the top of any endpoint that should be
    protected in pilot/production but open in demo mode.

    Uses the module-level `_MAINTENANCE_TOKEN` (read once at startup) so
    the value is consistent across the request lifetime and avoids repeated
    os.environ lookups.
    """
    if not _GUARD_DESTRUCTIVE:
        return
    if not _MAINTENANCE_TOKEN:
        raise HTTPException(
            status_code=503,
            detail=(
                "GUARD_DESTRUCTIVE=1 is set but MAINTENANCE_TOKEN is empty. "
                "Set MAINTENANCE_TOKEN to a strong secret to use guarded endpoints."
            ),
        )
    token = request.headers.get("X-Guard-Token", "")
    if not hmac.compare_digest(token.encode(), _MAINTENANCE_TOKEN.encode()):
        raise HTTPException(
            status_code=403,
            detail="Destructive operation requires a valid X-Guard-Token header.",
        )


# ── Email notification stub ───────────────────────────────────────────────────

def _send_notification_email(student_name: str, action: str, score: float) -> None:
    """Stub for SendGrid email notification. Replace with real implementation."""
    import logging
    log = logging.getLogger(__name__)
    log.info(
        "EMAIL NOTIFICATION [stub] → action=%s student=%s score=%.3f — "
        "integrate SendGrid here: https://docs.sendgrid.com/api-reference/mail-send/mail-send",
        action, student_name, score
    )
    # TODO: Replace with actual SendGrid call:
    # from sendgrid import SendGridAPIClient
    # from sendgrid.helpers.mail import Mail
    # sg = SendGridAPIClient(os.environ.get('SENDGRID_API_KEY'))
    # message = Mail(from_email='noreply@original.ai', to_emails=professor_email, ...)
    # sg.send(message)


# ── Health ────────────────────────────────────────────────────────────────────

@app.get("/health", response_model=HealthResponse)
def health():
    return HealthResponse(
        status="ok",
        feature_dim=FEATURE_DIM,
        students_in_store=store.count(),
    )


@app.get("/admin/health")
def admin_health():
    """
    System health summary for the admin dashboard.

    Returns student count, manifest totals, and queue depth from the live store.
    Latency is computed from the most recent manifest entries where available.
    """
    student_count = store.count()

    # Pull manifest stats for submission / flag counts
    try:
        stats = store.manifest_stats()
    except Exception:
        stats = {}

    total_submissions = stats.get("total", 0)
    flagged_count = stats.get("by_action", {}).get("escalate", 0) + \
                   stats.get("by_action", {}).get("schedule_conversation", 0)

    # Estimate avg latency from recent manifests (created_at timestamps)
    avg_latency_ms = None
    try:
        recent = store.list_manifests(limit=20)
        items = recent.get("items", [])
        if items:
            # Use latency stored in manifest if present, else report None
            latencies = [
                item.get("latency_ms") for item in items
                if item.get("latency_ms") is not None
            ]
            if latencies:
                avg_latency_ms = round(sum(latencies) / len(latencies))
    except Exception:
        pass

    return {
        "api_status": "operational",
        "student_count": student_count,
        "total_submissions": total_submissions,
        "flagged_count": flagged_count,
        "avg_latency_ms": avg_latency_ms,
        "queue_depth": 0,   # demo server processes synchronously; always 0
        "uptime_pct": 99.97,
    }


# ── Staff auth: email + password → principal token (ADR-003, Phase 1.x) ───────
# Professors / admins / operators log in here. Students use student_auth.
# Every method (this, and LTI later) mints the same principal token, which the
# tenant-isolation middleware then enforces. Demo needs no login — anonymous
# requests resolve to the demo principal and keep working.

@app.post("/auth/login")
def auth_login(body: dict):
    email = str(body.get("email") or "").strip()
    password = str(body.get("password") or "")
    if not email or not password:
        raise HTTPException(status_code=422, detail="email and password are required")
    user = users_mod.authenticate(email, password)
    if not user:
        _repo().log_audit(action="login", actor=email, result="denied")
        raise HTTPException(status_code=401, detail="Invalid email or password")
    token = principal_mod.mint_principal_token(
        user["user_id"], user["role"], user["tenant_id"]
    )
    _repo().log_audit(
        action="login", tenant_id=user["tenant_id"], actor=user["email"], result="ok"
    )
    return {
        "token": token,
        "role": user["role"],
        "tenant_id": user["tenant_id"],
        "name": user["name"],
        "email": user["email"],
    }


@app.get("/auth/me")
def auth_me(request: "Request"):
    """Return the authenticated principal, or 401 for anonymous/demo callers."""
    p = getattr(request.state, "principal", None)
    if p is None or p.is_demo:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return {
        "user_id": p.user_id,
        "role": p.role,
        "tenant_id": p.tenant_id,
        "auth_method": p.auth_method,
    }


@app.post("/auth/register", status_code=201)
def auth_register(body: dict, request: "Request"):
    """
    Provision a staff user. Privileged: guarded by GUARD_DESTRUCTIVE in
    pilot/production (X-Guard-Token required); open in demo for convenience.
    """
    _require_guard(request)
    email = str(body.get("email") or "").strip()
    password = str(body.get("password") or "")
    role = str(body.get("role") or "professor")
    tenant_id = str(body.get("tenant_id") or "").strip()
    name = str(body.get("name") or "")
    if not email or not password or not tenant_id:
        raise HTTPException(status_code=422, detail="email, password, and tenant_id are required")
    if role not in ("professor", "admin", "operator"):
        raise HTTPException(status_code=422, detail="role must be professor, admin, or operator")
    if len(password) < 8:
        raise HTTPException(status_code=422, detail="password must be at least 8 characters")
    if store.get_user_by_email(email):
        raise HTTPException(status_code=409, detail="a user with that email already exists")
    user = users_mod.create_user(email, password, role, tenant_id, name)
    _repo().log_audit(
        action="user_register", tenant_id=tenant_id, actor=email, result="ok",
        details={"role": role},
    )
    return {"user_id": user["user_id"], "email": user["email"], "role": role, "tenant_id": tenant_id}


# ── LTI 1.3 launch (ADR-003, Phase 1.5) ───────────────────────────────────────
# Lets an LMS (Canvas/Blackboard/Moodle) launch Original directly. The launch
# terminates in the same principal token as email/password login. Crypto deps
# are imported lazily, so the demo (which omits python-jose) still boots; the
# endpoints return a clear error until LTI is configured.

@app.api_route("/lti/login", methods=["GET", "POST"])
async def lti_login(request: "Request"):
    from . import lti
    params = dict(request.query_params)
    if request.method == "POST":
        form = await request.form()
        params.update({k: str(v) for k, v in form.items()})
    try:
        url = lti.build_login_redirect(params)
    except lti.LtiError as e:
        raise HTTPException(status_code=400, detail=f"LTI login error: {e}")
    return RedirectResponse(url, status_code=302)


@app.post("/lti/launch")
async def lti_launch(request: "Request"):
    from . import lti
    form = await request.form()
    id_token = str(form.get("id_token", ""))
    state = str(form.get("state", ""))
    if not id_token or not state:
        raise HTTPException(status_code=400, detail="missing id_token or state")
    try:
        claims = lti.verify_launch(id_token, state)
    except lti.LtiError as e:
        raise HTTPException(status_code=401, detail=f"LTI launch rejected: {e}")
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="LTI requires python-jose, which is not installed in this deployment.",
        )
    p = lti.principal_from_claims(claims)
    _repo().log_audit(
        action="lti_launch", tenant_id=p["tenant_id"],
        actor=str(claims.get("sub", "")), result="ok",
        details={"role": p["role"], "redirect": p.get("redirect")},
    )
    # All localStorage keys the destination needs (token + identity + any binding).
    ls = {
        p["token_key"]: p["token"],
        "original_role": p["role"],
        "original_tenant": p["tenant_id"],
    }
    ls.update(p.get("extra") or {})
    redirect = p.get("redirect") or "professor.html"
    params = p.get("params") or {}
    if params and redirect.endswith("/"):
        redirect = redirect + "?" + urllib.parse.urlencode(params)
    sets = "".join(
        f"localStorage.setItem({json.dumps(k)},{json.dumps(v)});" for k, v in ls.items()
    )
    # Hand the token to the browser (server-rendered, never in a URL) and break
    # out of the LMS iframe into the full-page app.
    html = (
        "<!doctype html><meta charset=utf-8><title>Bluebook · Original</title>"
        "<body style=\"font-family:Inter,system-ui;background:#001020;color:#C9A961;"
        "display:flex;align-items:center;justify-content:center;height:100vh;margin:0\">"
        "<div style=\"font-family:'Cormorant Garamond',Georgia,serif;font-size:1.3rem\">Entering examination…</div>"
        f"<script>try{{{sets}}}catch(e){{}}"
        f"var u={json.dumps(redirect)};try{{window.top.location.replace(u);}}catch(e){{location.replace(u);}}"
        "</script></body>"
    )
    return HTMLResponse(html)


@app.get("/lti/jwks")
def lti_jwks():
    from . import lti
    return lti.public_jwks()


# ── Bluebook examinations (secure-exam layer, tenant-scoped) ──────────────────
# Instructor-created exams persist here. Submissions themselves flow to
# /students/{id}/baseline as proctored samples. Scoping mirrors list_students:
# an authenticated non-super principal sees only its tenant; demo sees "demo".

def _bluebook_tenant(request: "Request") -> str:
    p = getattr(request.state, "principal", None)
    if p and not p.is_demo:
        return p.tenant_id
    return principal_mod.DEMO_TENANT


def _int_or(v, default):
    try:
        return int(v)
    except (TypeError, ValueError):
        return default


@app.post("/bluebook/exams", status_code=201)
def bluebook_create_exam(body: dict, request: "Request"):
    title = str(body.get("title") or "").strip()
    if not title:
        raise HTTPException(status_code=422, detail="title is required")
    tenant = _bluebook_tenant(request)
    rec = {
        "id": uuid.uuid4().hex[:16],
        "tenant_id": tenant,
        "title": title[:200],
        "course": str(body.get("course") or "")[:80],
        "duration": _int_or(body.get("duration"), 90),
        "minWords": _int_or(body.get("minWords"), 0),
        "maxWords": _int_or(body.get("maxWords"), 0),
        "prompt": str(body.get("prompt") or "")[:8000],
        "conditions": body.get("conditions") if isinstance(body.get("conditions"), dict) else {},
        "status": (str(body.get("status") or "DRAFT").upper())[:20],
    }
    store.put_bluebook_exam(rec)
    _repo().log_audit(action="bluebook_exam_create", tenant_id=tenant, details={"title": title})
    rec["submissions"] = 0
    return rec


@app.get("/bluebook/exams")
def bluebook_list_exams(request: "Request"):
    p = getattr(request.state, "principal", None)
    if p and not p.is_demo and p.role not in principal_mod.SUPER_ROLES:
        exams = store.list_bluebook_exams(p.tenant_id)
    elif p and p.is_demo:
        exams = store.list_bluebook_exams(principal_mod.DEMO_TENANT)
    else:  # super / operator → all tenants
        exams = store.list_bluebook_exams(None)
    return {"exams": exams}


@app.get("/bluebook/exams/{exam_id}")
def bluebook_get_exam(exam_id: str, request: "Request"):
    rec = store.get_bluebook_exam(exam_id)
    if not rec:
        raise HTTPException(status_code=404, detail="exam not found")
    p = getattr(request.state, "principal", None)
    owner = rec.get("tenant_id")
    if p and not p.is_demo and p.role not in principal_mod.SUPER_ROLES and owner != p.tenant_id:
        raise HTTPException(status_code=403, detail="cross-tenant access denied")
    if p and p.is_demo and owner not in (None, principal_mod.DEMO_TENANT):
        raise HTTPException(status_code=403, detail="cross-tenant access denied")
    return rec


@app.post("/bluebook/submissions", status_code=201)
def bluebook_record_submission(body: dict, request: "Request"):
    """Record one sat examination (the integrity reading for the Results view)."""
    tenant = _bluebook_tenant(request)

    def _clamp_pct(v):
        n = _int_or(v, None)
        return None if n is None else max(0, min(100, n))

    rec = {
        "id": uuid.uuid4().hex[:16],
        "exam_id": (str(body.get("exam_id")) if body.get("exam_id") else None),
        "tenant_id": tenant,
        "student_id": str(body.get("student_id") or "")[:128],
        "candidate": str(body.get("candidate") or "")[:120],
        "exam_title": str(body.get("exam_title") or "")[:200],
        "course": str(body.get("course") or "")[:80],
        "word_count": _int_or(body.get("word_count"), 0),
        "time_min": _int_or(body.get("time_min"), 0),
        "stylometric": _clamp_pct(body.get("stylometric")),
        "ai_score": _clamp_pct(body.get("ai_score")),
        "status": (str(body.get("status") or "SUBMITTED").upper())[:20],
    }
    store.put_bluebook_submission(rec)
    _repo().log_audit(action="bluebook_submission", tenant_id=tenant,
                      student_id=rec["student_id"], details={"exam_id": rec["exam_id"]})
    return {"id": rec["id"], "status": rec["status"]}


@app.get("/bluebook/submissions")
def bluebook_list_submissions(request: "Request"):
    p = getattr(request.state, "principal", None)
    if p and not p.is_demo and p.role not in principal_mod.SUPER_ROLES:
        subs = store.list_bluebook_submissions(p.tenant_id)
    elif p and p.is_demo:
        subs = store.list_bluebook_submissions(principal_mod.DEMO_TENANT)
    else:
        subs = store.list_bluebook_submissions(None)
    return {"submissions": subs}


@app.post("/bluebook/courses", status_code=201)
def bluebook_create_course(body: dict, request: "Request"):
    name = str(body.get("name") or "").strip()
    if not name:
        raise HTTPException(status_code=422, detail="course name is required")
    tenant = _bluebook_tenant(request)
    rec = {
        "id": uuid.uuid4().hex[:16],
        "tenant_id": tenant,
        "code": str(body.get("code") or "")[:40],
        "name": name[:160],
        "term": str(body.get("term") or "")[:60],
        "status": (str(body.get("status") or "ACTIVE").upper())[:20],
    }
    store.put_bluebook_course(rec)
    _repo().log_audit(action="bluebook_course_create", tenant_id=tenant, details={"code": rec["code"]})
    return {**rec, "active": rec["status"] == "ACTIVE", "students": 0, "exams": 0}


@app.get("/bluebook/courses")
def bluebook_list_courses(request: "Request"):
    p = getattr(request.state, "principal", None)
    if p and not p.is_demo and p.role not in principal_mod.SUPER_ROLES:
        courses = store.list_bluebook_courses(p.tenant_id)
    elif p and p.is_demo:
        courses = store.list_bluebook_courses(principal_mod.DEMO_TENANT)
    else:
        courses = store.list_bluebook_courses(None)
    return {"courses": courses}


# ── Student list ──────────────────────────────────────────────────────────────

@app.get("/students")
def list_students(request: "Request", tenant_id: str = ""):
    """
    List student IDs.

    Scoping (ADR-003): an authenticated, non-super principal only ever sees its
    own tenant's students — the `tenant_id` query param cannot widen that. The
    anonymous demo principal and super/operator roles keep the original
    behaviour (optional `tenant_id` filter, else all).
    """
    principal = getattr(request.state, "principal", None)
    if principal and not principal.is_demo and principal.role not in principal_mod.SUPER_ROLES:
        ids = store.list_ids_for_tenant(principal.tenant_id)
    elif tenant_id:
        ids = store.list_ids_for_tenant(tenant_id)
    else:
        ids = store.list_ids()
    return {"students": ids}


# ── Student state ─────────────────────────────────────────────────────────────

@app.get("/students/{student_id}", response_model=StudentStateResponse)
def get_student(student_id: str):
    state = store.get(student_id)
    if state is None:
        raise HTTPException(status_code=404, detail=f"Student '{student_id}' not found")

    traj = state.trajectory
    baseline_dict = {
        code: float(state.baseline_mean[i])
        for i, code in enumerate(
            __import__("original.constants", fromlist=["ALL_FEATURE_CODES"]).ALL_FEATURE_CODES
        )
    }

    samples_out = [
        SampleSummary(
            index=i,
            assignment=s.assignment,
            provenance=s.provenance,
            submitted_at=s.submitted_at,
            auth_weight=s.auth_weight,
        )
        for i, s in enumerate(state.samples)
    ]

    return StudentStateResponse(
        student_id=student_id,
        sample_count=state.sample_count,
        authenticated_count=state.authenticated_count,
        purity=state.purity,
        effective_sample_count=state.effective_sample_count,
        trajectory_direction=traj.direction,
        trajectory_confidence=traj.confidence,
        baseline_vector=baseline_dict,
        samples=samples_out,
    )


# ── Read a single baseline sample's prose text ───────────────────────────────
# The StudentStateResponse exposes only SampleSummary metadata (index,
# assignment, provenance, submitted_at, auth_weight) — not the text itself,
# so the demo UI can stay slim. When the professor wants to read a specific
# sample's writing (to remind themselves of the student's voice, or to
# verify a sample is legitimate before authenticating it), they fetch
# the prose lazily via this endpoint.

@app.get("/students/{student_id}/samples/{index}/text")
def get_sample_text(student_id: str, index: int):
    """
    Return the raw prose of a single baseline sample.

    Returns 404 if the student doesn't exist or the index is out of range.
    Response shape mirrors the SampleSummary metadata so the caller can
    render headers + body in a single round-trip without re-fetching the
    student state.
    """
    state = store.get(student_id)
    if state is None:
        raise HTTPException(status_code=404, detail=f"Student '{student_id}' not found")
    if index < 0 or index >= len(state.samples):
        raise HTTPException(
            status_code=404,
            detail=f"Sample index {index} out of range (student has {len(state.samples)} samples)",
        )
    s = state.samples[index]
    return {
        "student_id": student_id,
        "index": index,
        "assignment": s.assignment,
        "provenance": s.provenance,
        "submitted_at": s.submitted_at,
        "auth_weight": s.auth_weight,
        "word_count": len((s.text or "").split()),
        "char_count": len(s.text or ""),
        "text": s.text or "",
    }


# ── FERPA: student data deletion ──────────────────────────────────────────────

@app.delete("/students/{student_id}", status_code=200)
def delete_student(student_id: str, request: "Request"):
    """
    Permanently delete all stored data for a student (FERPA right-to-erasure).

    Removes the student profile, all baseline samples, all fidelity scores,
    all adaptive-context manifests, and all instructor corrections associated
    with this student_id.  The deletion is immediate and irreversible — there
    is no soft-delete or recovery path.

    Returns 200 with a confirmation payload on success, 404 if not found.
    Returns 404 also when the SQLite commit fails (no data was removed).

    Intended audience: institution data-compliance officers and LMS admins.
    When GUARD_DESTRUCTIVE=1 (pilot/production mode), requires an
    X-Guard-Token header matching MAINTENANCE_TOKEN. Demo mode is open.
    """
    _require_guard(request)
    remote = getattr(request.client, "host", "unknown") if request.client else "unknown"
    deleted = store.delete_student(student_id)
    if not deleted:
        _repo().log_audit(action="student_delete", student_id=student_id, actor=remote, result="not_found")
        raise HTTPException(
            status_code=404,
            detail=f"Student '{student_id}' not found — nothing to delete.",
        )
    _repo().log_audit(action="student_delete", student_id=student_id, actor=remote, result="ok")
    return {
        "deleted": True,
        "student_id": student_id,
        "message": (
            f"All data for student '{student_id}' has been permanently removed "
            "(baseline profile, fidelity scores, manifests, corrections)."
        ),
    }


# ── Tenant registry ───────────────────────────────────────────────────────────
# Phase 0 foundation: lightweight per-institution metadata stored in SQLite.
# Lets demo operator register schools with an environment label (demo/pilot/
# production) before Postgres multi-tenancy is needed.

@app.post("/tenants", status_code=201)
def create_tenant(body: dict):
    """
    Register or update a tenant (institution) record.

    Required body fields:
        tenant_id   — stable slug (e.g. 'seminary-of-dallas')
        name        — human-readable institution name

    Optional body fields:
        environment — 'demo' | 'pilot' | 'production'  (default: 'demo')
        meta        — arbitrary dict of metadata (contact email, LMS URL, etc.)
                      Capped at 10 keys, values must be strings ≤ 500 chars.
    """
    tenant_id = str(body.get("tenant_id", "")).strip()
    name = str(body.get("name", "")).strip()
    if not tenant_id or not name:
        raise HTTPException(status_code=422, detail="tenant_id and name are required")
    if len(tenant_id) > 80 or len(name) > 200:
        raise HTTPException(status_code=422, detail="tenant_id max 80 chars, name max 200 chars")
    environment = body.get("environment", "demo")
    if environment not in ("demo", "pilot", "production"):
        raise HTTPException(status_code=422, detail="environment must be 'demo', 'pilot', or 'production'")
    # Validate meta payload — prevents unbounded JSON storage
    meta = body.get("meta") or {}
    if not isinstance(meta, dict):
        raise HTTPException(status_code=422, detail="meta must be a JSON object")
    if len(meta) > 10:
        raise HTTPException(status_code=422, detail="meta must have at most 10 keys")
    meta = {str(k)[:80]: str(v)[:500] for k, v in list(meta.items())[:10]}
    _repo().put_tenant(tenant_id, name, environment=environment, meta=meta)
    principal_mod.invalidate_tenant_cache()  # env may have changed → drop stale cache
    _repo().log_audit(action="tenant_register", tenant_id=tenant_id, details={"name": name, "environment": environment})
    return {"tenant_id": tenant_id, "name": name, "environment": environment}


@app.get("/tenants")
def list_tenants(environment: str = ""):
    """List all registered tenants, optionally filtered by environment."""
    return _repo().list_tenants(environment=environment or None)


@app.get("/tenants/{tenant_id}")
def get_tenant(tenant_id: str):
    """Get a single tenant record."""
    t = _repo().get_tenant(tenant_id)
    if not t:
        raise HTTPException(status_code=404, detail=f"Tenant '{tenant_id}' not found")
    return t


@app.get("/tenants/{tenant_id}/stats")
def tenant_stats(tenant_id: str):
    """
    Aggregate statistics for a tenant — student count, submission volume,
    action breakdown, last active timestamp.

    Used by the operator dashboard to show all-schools-at-a-glance.
    """
    t = _repo().get_tenant(tenant_id)
    if not t:
        raise HTTPException(status_code=404, detail=f"Tenant '{tenant_id}' not found")
    return _repo().tenant_stats(tenant_id)


@app.delete("/tenants/{tenant_id}/students", status_code=200)
def delete_tenant_students(tenant_id: str, request: "Request"):
    """
    FERPA-safe bulk deletion of all students belonging to a tenant.

    Iterates list_ids_for_tenant() and calls store.delete_student() for each —
    the same code path as individual deletion, so all linked records
    (fidelity scores, manifests, corrections, audit rows) are purged.

    When GUARD_DESTRUCTIVE=1, requires X-Guard-Token header.
    """
    _require_guard(request)
    t = _repo().get_tenant(tenant_id)
    if not t:
        raise HTTPException(status_code=404, detail=f"Tenant '{tenant_id}' not found")
    result = store.delete_tenant_students(tenant_id)
    return {
        "tenant_id":    tenant_id,
        "deleted_count": result["deleted_count"],
        "failed_ids":   result["failed_ids"],
        "message": (
            f"Deleted {result['deleted_count']} student(s) from '{tenant_id}'. "
            + (f"Failed: {result['failed_ids']}" if result["failed_ids"] else "")
        ).strip(),
    }


# ── Student authentication (converged path) ──────────────────────────────────
# A student signs in with (institution, email). Their id is derived
# deterministically (institution-scoped email hash), their institution is
# auto-registered as a demo tenant, and they receive a signed, stateless
# session token. No password in the demo path — identity is the email +
# institution, which the v1 path can later harden with a real credential.

@app.post("/student-auth/login")
def student_login(body: dict, request: "Request"):
    """
    Sign a student in. Body: { email, institution, name? }.

    Derives an institution-scoped student id, ensures the institution exists in
    the tenant registry (auto-provisioned as a demo tenant), creates the
    student record if new, and returns a signed session token.
    """
    email = str(body.get("email") or "").strip()
    institution = str(body.get("institution") or "").strip()
    name = str(body.get("name") or "").strip()
    if not email or "@" not in email:
        raise HTTPException(status_code=422, detail="A valid email is required.")
    if not institution:
        raise HTTPException(status_code=422, detail="An institution is required.")

    tenant_id = student_auth.slugify(institution)
    student_id = student_auth.derive_student_id(institution, email)

    # Auto-provision the institution as a demo tenant (idempotent).
    if not _repo().get_tenant(tenant_id):
        _repo().put_tenant(tenant_id, institution, environment="demo",
                           meta={"auto_provisioned": "student_login"})

    # Ensure the student record exists so the dashboard has somewhere to read.
    store.get_or_create(student_id)

    token = student_auth.mint_session(student_id, name or email.split("@")[0])
    remote = getattr(request.client, "host", "unknown") if request.client else "unknown"
    _repo().log_audit(action="student_login", student_id=student_id,
                      tenant_id=tenant_id, actor=remote)
    return {
        "token":       token,
        "student_id":  student_id,
        "name":        name or email.split("@")[0],
        "tenant_id":   tenant_id,
        "institution": institution,
    }


@app.get("/student-auth/me")
def student_me(request: "Request"):
    """
    Resolve the current student from the session token (Authorization: Bearer
    <token> or X-Student-Token header). 401 if missing/invalid/expired.
    """
    auth = request.headers.get("Authorization", "")
    token = auth[7:] if auth.lower().startswith("bearer ") else request.headers.get("X-Student-Token", "")
    session = student_auth.verify_session(token)
    if not session:
        raise HTTPException(status_code=401, detail="Not signed in.")
    return {"student_id": session["sid"], "name": session.get("name", "")}


# ── FERPA: data inventory + audit log ────────────────────────────────────────

@app.get("/students/{student_id}/data-inventory")
def student_data_inventory(student_id: str):
    """
    FERPA data-access response: structured inventory of all data held for a student.

    Returns a categorized breakdown of:
    - Baseline writing samples (count, provenance types, date range)
    - Fidelity / calibration scores
    - Scored submission manifests (by recommendation action)
    - Instructor corrections
    - Audit log entries

    Intended for: student data-access requests, FERPA compliance officers,
    deletion confirmations ("prove everything was purged").
    """
    inv = store.student_data_inventory(student_id)
    if inv is None:
        raise HTTPException(status_code=404, detail=f"Student '{student_id}' not found")
    return inv


@app.get("/admin/audit")
def list_audit_log(
    student_id: str = "",
    action: str = "",
    limit: int = 100,
    offset: int = 0,
):
    """
    Query the system audit log.

    Optional filters:
        student_id — restrict to a specific student
        action     — restrict to a specific action type
                     (baseline_add, score, student_delete, correction,
                      threshold_apply, tenant_register, bulk_delete)

    Results are ordered most-recent-first.
    """
    limit = min(limit, 500)
    return _repo().list_audit(
        student_id=student_id or None,
        action=action or None,
        limit=limit,
        offset=offset,
    )


# ── Add baseline sample ───────────────────────────────────────────────────────

@app.post("/students/{student_id}/baseline")
def add_baseline(student_id: str, req: AddSampleRequest):
    if req.provenance not in AUTH_WEIGHTS:
        raise HTTPException(
            status_code=422,
            detail=f"provenance must be one of: {list(AUTH_WEIGHTS)}"
        )

    state = store.get_or_create(student_id)
    vec = feature_vector(req.text, keystroke_data=req.keystroke_data)

    # Genre label — classify the text at ingestion time so the Hierarchical
    # Bayesian prior (BAYESIAN_PRIOR_ENABLED=1) has cross-student genre data.
    # Uses the same rule-based resolver as the context manifest pipeline.
    # Runs even when the manifest flag is off — genre metadata is cheap and
    # the prior needs it independent of the manifest subsystem.
    _sample_genre: Optional[str] = None
    try:
        from .context.resolvers import resolve_genre
        _genre_result = resolve_genre(req.text)
        _sample_genre = (_genre_result or {}).get("primary")
    except Exception:
        pass   # genre labeling is best-effort; don't fail baseline ingestion

    sample = BaselineSample(
        text=req.text,
        vector=vec,
        provenance=req.provenance,
        auth_weight=AUTH_WEIGHTS[req.provenance],
        assignment=req.assignment,
        submitted_at=req.submitted_at,
        genre=_sample_genre,
    )

    # ── Phase 8: drift gate before adding to baseline ─────────────────────────
    # Only authenticated samples (auth_weight > 0) participate in the
    # baseline_mean — unverified samples can't drift the baseline either way,
    # so we skip the check for them. The check is best-effort: a failure is
    # logged and the sample is admitted as before (Phase 1 behaviour).
    drift_result = None
    if AUTH_WEIGHTS[req.provenance] > 0:
        try:
            drift_result = state.check_drift(sample)
        except Exception as e:
            logging.getLogger(__name__).warning(
                "drift check failed for %s: %s — admitting sample without gate",
                student_id, e,
            )
            drift_result = None

    # check_drift mutates _consecutive_drift_count regardless of recommendation;
    # persist the counter even on flag/rebaseline so the workflow is sticky.
    if drift_result is not None and drift_result.recommendation != "accept":
        # Sample is held for review — DO NOT admit to state.samples.
        store.put(state)   # persist counter mutation
        body = DriftPendingResponse(
            status="pending_review" if drift_result.recommendation == "flag_for_review"
                   else "rebaseline_required",
            student_id=student_id,
            drift=DriftResultOut(**drift_result.to_dict()),
        )
        # 202 = Accepted but not applied (review pending);
        # 409 = Conflict (existing baseline is stale, rebaseline needed).
        status_code = 202 if drift_result.recommendation == "flag_for_review" else 409
        raise HTTPException(status_code=status_code, detail=body.model_dump())

    state.add_sample(sample)

    # Update tension arc κ baseline for authenticated samples
    if req.provenance in ("proctored", "verified"):
        arc = analyze_tension_arc(req.text)
        if arc.catastrophe_index > 0:   # skip insufficient-length samples
            new_mean = update_student_baseline_kappa(state.kappa_log, arc.catastrophe_index)
            state.baseline_kappa = new_mean

    store.put(state)   # persist to SQLite

    # Audit log — record the baseline addition
    _repo().log_audit(
        action="baseline_add",
        student_id=student_id,
        details={
            "provenance":         req.provenance,
            "auth_weight":        AUTH_WEIGHTS[req.provenance],
            "sample_count_after": state.sample_count,
            "genre":              _sample_genre,
        },
    )

    # Auto-complete any outstanding magic-link baseline requests for this
    # student (Phase 2). Only fires for authenticated provenance — an
    # unverified self-upload doesn't satisfy a "proctored baseline" request.
    completed_requests: list = []
    if AUTH_WEIGHTS[req.provenance] > 0:
        try:
            completed_requests = baseline_requests.mark_completed_for_student(student_id)
        except Exception as e:
            logging.getLogger(__name__).warning(
                "baseline-request auto-complete failed for %s: %s", student_id, e,
            )

    response = {
        "student_id": student_id,
        "sample_index": state.sample_count - 1,
        "provenance": req.provenance,
        "auth_weight": AUTH_WEIGHTS[req.provenance],
        "authenticated_count": state.authenticated_count,
        "purity": state.purity,
    }
    # Include the drift result on accept too — useful for UIs that want to
    # show the trend even when no action was triggered.
    if drift_result is not None:
        response["drift"] = drift_result.to_dict()
    if completed_requests:
        response["completed_baseline_requests"] = [
            r.external_request_id for r in completed_requests
        ]
    return response


# ── Bbook integration: request a proctored baseline sitting ──────────────────
# Phase 2 (Original-first flow). The professor on professor.html clicks
# "Request proctored baseline" for a student. Original calls Bbook to
# provision a one-off magic-link exam and records the pending request here
# so the professor can see status. When Bbook later POSTs the resulting
# baseline back to /students/{id}/baseline (Phase 1 sync flow), the
# corresponding pending request is auto-marked completed.

from pydantic import BaseModel as _PydanticBaseModel  # local import to avoid disturbing top imports


class RequestBaselineRequest(_PydanticBaseModel):
    """Inbound shape for POST /students/{id}/request-baseline."""
    student_email: str
    student_name: str
    exam_title: str = "Proctored Baseline Sitting"
    institution_name: Optional[str] = None
    requested_by: Optional[str] = None    # free-form audit field
    duration_mins: int = 45
    min_word_count: Optional[int] = None
    max_word_count: Optional[int] = None
    prompt_text: Optional[str] = None


@app.post("/students/{student_id}/request-baseline")
def request_proctored_baseline(student_id: str, req: RequestBaselineRequest):
    """
    Provision a magic-link proctored baseline exam in Bbook for this student.

    Returns the pending request record with the magic-link URL (only when
    SMTP delivery failed or is unconfigured — otherwise the student receives
    it by email). Idempotency is per-call: each invocation creates a new
    pending request with a fresh UUID.

    Requires BBOOK_API_URL and BBOOK_EXTERNAL_SECRET in the environment.
    Returns 503 if Bbook integration is not configured, 502 on Bbook errors.
    """
    if not bbook_client.is_enabled():
        raise HTTPException(
            status_code=503,
            detail="Bbook integration is not configured (set BBOOK_API_URL).",
        )

    external_id = baseline_requests.make_external_id()

    # Pre-record the pending request so the UI sees it immediately, even
    # before the Bbook round-trip completes. We'll update with the magic
    # link and Bbook exam id once the response arrives.
    import time as _time
    pending = baseline_requests.BaselineRequest(
        external_request_id=external_id,
        student_id=student_id,
        student_email=req.student_email,
        student_name=req.student_name,
        exam_title=req.exam_title,
        bbook_exam_id=None,
        magic_link=None,
        requested_at=_time.time(),
        expires_at=None,
        requested_by=req.requested_by,
    )
    baseline_requests.record(pending)

    try:
        result = bbook_client.request_baseline(
            student_email=req.student_email,
            student_name=req.student_name,
            exam_title=req.exam_title,
            institution_name=req.institution_name,
            requested_by=req.requested_by,
            duration_mins=req.duration_mins,
            min_word_count=req.min_word_count,
            max_word_count=req.max_word_count,
            prompt_text=req.prompt_text,
            external_request_id=external_id,
        )
    except Exception as e:
        baseline_requests.mark_failed(external_id, str(e))
        logging.getLogger(__name__).exception("Bbook baseline-request call failed")
        raise HTTPException(status_code=502, detail=f"Bbook call failed: {e}")

    # Update the pending record with the Bbook exam id + magic link + expiry.
    pending.bbook_exam_id = result.examId
    pending.magic_link = result.magicLink
    pending.email_delivered = result.emailDelivered
    if result.expiresAt:
        # Parse "2026-05-18T..." to epoch seconds for the registry
        from datetime import datetime
        try:
            pending.expires_at = datetime.fromisoformat(
                result.expiresAt.replace("Z", "+00:00")
            ).timestamp()
        except Exception:
            pending.expires_at = None
    baseline_requests.record(pending)

    return pending.to_dict()


@app.get("/baseline-requests/pending")
def list_pending_baseline_requests():
    """List all currently-pending proctored baseline requests."""
    return {"requests": [r.to_dict() for r in baseline_requests.list_pending()]}


@app.get("/baseline-requests")
def list_all_baseline_requests(request: "Request"):
    """
    List every proctored baseline request, regardless of status.
    When GUARD_DESTRUCTIVE=1, requires X-Guard-Token header (admin only).
    """
    _require_guard(request)
    return {"requests": [r.to_dict() for r in baseline_requests.list_all()]}


# ── Formation pathways (ADR-002 — routed through the Repository seam) ─────────
# These handlers depend only on the Repository interface, never on store
# directly. Swapping in a Postgres-backed Repository requires no change here.

@app.get("/students/{student_id}/formation")
def get_formation(student_id: str):
    """Return the student's active (or most recent) formation pathway, or null."""
    repo = get_repository(os.environ.get("ENVIRONMENT", "demo"))
    return {"pathway": repo.get_formation_pathway(student_id)}


@app.post("/students/{student_id}/formation", status_code=201)
def open_formation(student_id: str, body: Optional[dict] = None):
    """
    Open a three-session formation pathway. Idempotent — returns the existing
    open pathway if one is already in progress.

    Optional body: { submission_id, reason }
    """
    body = body or {}
    repo = get_repository(os.environ.get("ENVIRONMENT", "demo"))
    pathway = repo.open_formation_pathway(
        student_id,
        submission_id=body.get("submission_id"),
        reason=body.get("reason"),
    )
    if pathway is None:
        raise HTTPException(status_code=500, detail="Could not open formation pathway")
    return {"pathway": pathway}


@app.post("/students/{student_id}/formation/advance")
def advance_formation(student_id: str):
    """
    Advance the open pathway by one session. On the final session the pathway
    completes and the triggering submission's review flag is cleared.
    Returns 404 if there is no open pathway.
    """
    repo = get_repository(os.environ.get("ENVIRONMENT", "demo"))
    pathway = repo.advance_formation_pathway(student_id)
    if pathway is None:
        raise HTTPException(
            status_code=404,
            detail=f"No open formation pathway for student '{student_id}'.",
        )
    return {"pathway": pathway}


# ── File upload (text extraction) ────────────────────────────────────────────

@app.post("/students/{student_id}/upload")
async def upload_file(student_id: str, file: UploadFile = File(...)):
    """Extract plain text from an uploaded .txt, .docx, or .pdf file."""
    filename = file.filename or "unknown"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    raw = await file.read()

    if ext == "txt":
        text = raw.decode("utf-8", errors="replace")
    elif ext == "docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx not installed")
    elif ext == "pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            text = "\n\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="pypdf not installed")
    else:
        raise HTTPException(
            status_code=422,
            detail=f"Unsupported file type '.{ext}'. Use .txt, .docx, or .pdf.",
        )

    word_count = len(text.split())
    return {"text": text, "filename": filename, "word_count": word_count}


# ── Score submission ──────────────────────────────────────────────────────────

@app.post("/students/{student_id}/score", response_model=Layer7OutputResponse)
def score_submission(student_id: str, req: ScoreSubmissionRequest, force: bool = False):
    state = store.get(student_id)
    if state is None:
        raise HTTPException(
            status_code=404,
            detail=f"Student '{student_id}' not found. Add baseline samples first."
        )
    if state.authenticated_count == 0:
        raise HTTPException(
            status_code=422,
            detail="No authenticated baseline samples found. "
                   "Add at least one 'proctored' or 'verified' sample first."
        )

    # Check cache only if force is False (allow cache bypass with force=True)
    submission_id = req.submission_id or f"{student_id}_submission_{state.sample_count}"
    if not force:
        # Check for cached result (stub for future caching implementation)
        existing_result = None  # TODO: retrieve from cache by submission_id
        if existing_result:
            return _to_response(existing_result)

    # ── Phase 5: adaptive-context orchestrator (env-flag gated) ───────────────
    # When both CONTEXT_MANIFEST_ENABLED and ADAPTIVE_WEIGHTS_ENABLED are
    # unset, the orchestrator short-circuits to plain extract_features +
    # feature_vector, preserving Phase 1 byte-identical behaviour.
    enable_manifest  = os.environ.get("CONTEXT_MANIFEST_ENABLED") == "1"
    enable_adaptive  = os.environ.get("ADAPTIVE_WEIGHTS_ENABLED") == "1"

    try:
        from .context.pipeline import run_adaptive_pipeline
        adaptive = run_adaptive_pipeline(
            text=req.text,
            state=state,
            submission_id=submission_id,
            keystroke_data=req.keystroke_data,
            enable_manifest=enable_manifest,
            enable_adaptive_weights=enable_adaptive,
        )
        feat_dict = adaptive.feat_dict
        vec       = adaptive.vector
        manifest  = adaptive.manifest
        adaptive_weights = adaptive.adaptive_weights
    except Exception as e:
        # Catastrophic orchestrator failure → fall through to the legacy path.
        # This guarantees that nothing in the new context layer can take down
        # the scoring endpoint, no matter how broken a resolver gets.
        logging.getLogger(__name__).warning(
            "Adaptive pipeline failed for %s: %s — falling back to Phase 1",
            submission_id, e,
        )
        feat_dict = extract_features(req.text, keystroke_data=req.keystroke_data)
        vec       = feature_vector(req.text, keystroke_data=req.keystroke_data)
        manifest  = None
        adaptive_weights = None

    manifest_dict = manifest.to_dict() if manifest is not None else None
    # n_tokens: thread the actual word count into the scorer so the Gaussian
    # wave packet attenuation in encode_amplitudes is proportional to the
    # real submission length, not a fixed default.
    _n_tokens = len(req.text.split())
    result = quantum_score(
        state=state,
        submission_vector=vec,
        feature_dict=feat_dict,
        submission_id=submission_id,
        adaptive_weights=adaptive_weights,
        manifest=manifest_dict,
        n_tokens=_n_tokens,
    )

    # ── Persist quantum fidelity for conformal calibration ───────────────────
    # Stores every scored fidelity so get_authentic_fidelities() can build
    # a calibration set for the conformal p-value on future submissions.
    # "Authentic" is approximated as action == no_action here; the instructor
    # corrections flow (put_correction + is_correct=True) should override
    # this for any verdict the professor marks as wrong.
    if result.authorship.quantum_fidelity > 0:
        try:
            store.put_fidelity_score(
                submission_id=submission_id,
                student_id=student_id,
                fidelity=result.authorship.quantum_fidelity,
                is_authentic=(result.recommendation.action == "no_action"),
            )
        except Exception as _e:
            logging.getLogger(__name__).debug(
                "put_fidelity_score skipped for %s: %s", submission_id, _e,
            )

    # ── Persist manifest to audit log when one was built ──────────────────────
    if manifest is not None:
        try:
            store.put_manifest(
                submission_id=submission_id,
                student_id=student_id,
                manifest=manifest,
                divergence_score=result.authorship.deviation_score,
                action=result.recommendation.action,
            )
        except Exception as e:
            logging.getLogger(__name__).warning(
                "Manifest audit-log write failed for %s: %s", submission_id, e,
            )

    # ── Phase 6: human-readable audit report (only when manifest exists) ──────
    # Built from the same triplet that drove the score: Layer7Output (math),
    # ContextManifest (directives), StudentState (sample provenance). When
    # there is no manifest (flag off), no report is produced — response stays
    # byte-identical to Phase 1.
    report = None
    if manifest is not None:
        try:
            from .context.report import build_report
            report = build_report(result, manifest, state)
        except Exception as e:
            logging.getLogger(__name__).warning(
                "Report assembly failed for %s: %s", submission_id, e,
            )

    # ── Tension Arc (runs alongside quantum score, independent signal) ────────
    arc = analyze_tension_arc(req.text, baseline_kappa=state.baseline_kappa)

    # ── Email notification stub for escalate/schedule_conversation actions ────
    action = result.recommendation.action
    overall_score = result.authorship.authorship_probability
    if action in ("escalate", "schedule_conversation"):
        _send_notification_email(student_name=student_id, action=action, score=overall_score)

    # ── Audit log — best-effort, never raises ─────────────────────────────────
    try:
        _repo().log_audit(
            action="score",
            student_id=student_id,
            details={
                "submission_id":   submission_id,
                "deviation_score": round(result.authorship.deviation_score, 4),
                "recommendation":  action,
                "sample_count":    state.sample_count,
            },
        )
    except Exception:
        pass

    return _to_response(result, arc, report=report)


# ── Serialisation helper ──────────────────────────────────────────────────────

def _to_response(r, arc=None, report=None) -> Layer7OutputResponse:
    """Convert internal dataclasses → Pydantic response model."""
    from .quantum.scoring import (
        Layer7Output, FeatureContribution, EntanglementAnomaly,
    )
    from .explainer import explain

    # Phase 6: ScoringReport → ScoringReportOut. Built upstream when a
    # manifest exists; None preserves Phase 1 byte-identical responses.
    report_out: Optional[ScoringReportOut] = None
    if report is not None:
        report_out = ScoringReportOut(**report.to_dict())

    return Layer7OutputResponse(
        student_id=r.student_id,
        submission_id=r.submission_id,
        authorship=AuthorshipSignalOut(
            authorship_probability=r.authorship.authorship_probability,
            deviation_score=r.authorship.deviation_score,
        ),
        trajectory=TrajectoryConformanceOut(
            direction=r.trajectory.direction,
            alignment=r.trajectory.alignment,
            confidence=r.trajectory.confidence,
            adjustment_factor=r.trajectory.adjustment_factor,
        ),
        interference=InterferenceDecompositionOut(
            total_probability=r.interference.total_probability,
            constructive_features=[
                FeatureContributionOut(**fc.__dict__)
                for fc in r.interference.constructive_features
            ],
            destructive_features=[
                FeatureContributionOut(**fc.__dict__)
                for fc in r.interference.destructive_features
            ],
            broken_entanglements=[
                EntanglementAnomalyOut(
                    feature_a=e.feature_a,
                    feature_b=e.feature_b,
                    tier_a=e.tier_a,
                    tier_b=e.tier_b,
                    anomaly_score=e.anomaly_score,
                    label=e.label,
                )
                for e in r.interference.broken_entanglements
            ],
            tier_breakdown=r.interference.tier_breakdown,
        ),
        baseline_confidence=BaselineConfidenceOut(
            purity=r.baseline_confidence.purity,
            sample_count=r.baseline_confidence.sample_count,
            authenticated_count=r.baseline_confidence.authenticated_count,
            effective_sample_count=r.baseline_confidence.effective_sample_count,
            trajectory_confidence=r.baseline_confidence.trajectory_confidence,
        ),
        domain=DomainSignalOut(
            theological_register_score=r.domain.theological_register_score,
            register_anomaly=r.domain.register_anomaly,
            confessional_balance=r.domain.confessional_balance,
        ),
        recommendation=RecommendedActionOut(
            action=r.recommendation.action,
            confidence=r.recommendation.confidence,
            rationale=r.recommendation.rationale,
        ),
        tension_arc=TensionArcOut(
            catastrophe_index=arc.catastrophe_index,
            resolution_ratio_mean=arc.resolution_ratio_mean,
            resolution_ratio_std=arc.resolution_ratio_std,
            mean_tension=arc.mean_tension,
            max_tension=arc.max_tension,
            authenticity_signal=arc.authenticity_signal,
            arc_flag=arc.arc_flag,
            arc_flag_reason=arc.arc_flag_reason,
            tension_series=arc.tension_series,
        ) if arc is not None else None,
        feature_vector=r.feature_vector,
        baseline_vector=r.baseline_vector,
        catastrophic_drift=getattr(r, 'catastrophic_drift', False),
        catastrophic_drift_rms_z=getattr(r, 'catastrophic_drift_rms_z', 0.0),
        # Phase 3: ContextManifestOut when CONTEXT_MANIFEST_ENABLED=1, else None.
        context_manifest=(
            ContextManifestOut(**getattr(r, 'context_manifest', None))
            if getattr(r, 'context_manifest', None) is not None
            else None
        ),
        # Phase 6: ScoringReportOut when a manifest+report were built.
        report=report_out,
        # Human-friendly explanation for professors/instructors
        human_explanation=explain(r),
    )


# ── Score audit log (best-effort, never raises) ───────────────────────────────
# Wire audit logging after the return object is built so any exception here
# cannot corrupt the response. The try/except is intentional insurance.


# ── Phase 7: sliding-window blend detection ──────────────────────────────────

@app.post(
    "/students/{student_id}/score/blend",
    response_model=BlendResultOut,
)
def score_blend(student_id: str, req: BlendDetectionRequest):
    """
    Detect mid-document fingerprint shifts (collaboration / AI insertion /
    advisor edits) by scoring overlapping token windows separately.

    Cost is N× the regular `/score` endpoint (one full feature extraction
    per window) — kept on a separate route so callers opt in explicitly.
    """
    state = store.get(student_id)
    if state is None:
        raise HTTPException(
            status_code=404,
            detail=f"Student '{student_id}' not found. Add baseline samples first.",
        )
    if state.authenticated_count == 0:
        raise HTTPException(
            status_code=422,
            detail="No authenticated baseline samples found. "
                   "Add at least one 'proctored' or 'verified' sample first.",
        )

    from .context.blend import detect_blend
    submission_id = req.submission_id or f"{student_id}_blend_{state.sample_count}"
    result = detect_blend(
        text=req.text,
        state=state,
        window_tokens=req.window_tokens,
        overlap=req.overlap,
        submission_id=submission_id,
    )
    return BlendResultOut(
        blend_detected=result.blend_detected,
        blend_index=result.blend_index,
        shift_positions=list(result.shift_positions),
        per_section=[
            WindowScoreOut(start=w.start, end=w.end,
                            score=w.score, confidence=w.confidence)
            for w in result.per_section
        ],
        n_tokens=result.n_tokens,
        fallback_reason=result.fallback_reason,
    )


# ── Batch file upload → baseline ──────────────────────────────────────────────

@app.post("/students/{student_id}/baseline/upload-batch")
async def upload_baseline_batch(
    student_id: str,
    files: List[UploadFile] = File(...),
    provenance: str = Form("verified"),
    assignment: str = Form(""),
):
    """
    Upload one or more files (PDF, DOCX, TXT) as baseline samples in a single
    request.  Mirrors the v1 batch upload but requires no auth — used by the
    Import Papers drawer in the professor demo.
    """
    if provenance not in AUTH_WEIGHTS:
        raise HTTPException(status_code=422, detail=f"provenance must be one of: {list(AUTH_WEIGHTS)}")

    state = store.get_or_create(student_id)
    imported = 0
    skipped_duplicates = 0
    errors: list[str] = []
    # Phase 8: per-file drift outcomes — surfaced on the batch response so
    # an instructor can see which files were held without aborting the batch.
    drift_holds: list[dict] = []

    for upload in files:
        filename = upload.filename or "unknown"
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        raw = await upload.read()

        # ── Text extraction ───────────────────────────────────────────────────
        try:
            if ext == "txt":
                text = raw.decode("utf-8", errors="replace")
            elif ext == "docx":
                from docx import Document as _Doc
                doc = _Doc(io.BytesIO(raw))
                text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            elif ext == "pdf":
                from pypdf import PdfReader as _PdfReader
                reader = _PdfReader(io.BytesIO(raw))
                text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
            else:
                errors.append(f"{filename}: unsupported type '.{ext}' — use .txt, .docx, or .pdf")
                continue
        except Exception as exc:
            errors.append(f"{filename}: extraction error — {exc}")
            continue

        if not text.strip():
            errors.append(f"{filename}: no text extracted (empty or image-only file?)")
            continue

        # ── Deduplication ─────────────────────────────────────────────────────
        import hashlib as _hashlib
        text_hash = _hashlib.sha256(text.encode()).hexdigest()
        if any(getattr(s, "text_hash", None) == text_hash for s in state.samples):
            skipped_duplicates += 1
            continue

        # ── Feature extraction & store ────────────────────────────────────────
        try:
            vec = feature_vector(text)
        except Exception as exc:
            errors.append(f"{filename}: feature extraction failed — {exc}")
            continue

        label = assignment.strip() or filename.rsplit(".", 1)[0]
        sample = BaselineSample(
            text=text,
            vector=vec,
            provenance=provenance,
            auth_weight=AUTH_WEIGHTS[provenance],
            assignment=label,
            submitted_at="",
        )
        # Attach hash for future dedup checks
        sample.text_hash = text_hash  # type: ignore[attr-defined]

        # ── Phase 8: per-file drift gate (best-effort) ────────────────────────
        # Batch ingestion does NOT 202/409 on drift — that would block the
        # whole upload. Instead we hold individual outliers, record them in
        # `drift_holds`, and continue the loop. Instructor sees the per-file
        # outcome in the response.
        if AUTH_WEIGHTS[provenance] > 0:
            try:
                dr = state.check_drift(sample)
                if dr.recommendation != "accept":
                    drift_holds.append({
                        "filename": filename,
                        "drift": dr.to_dict(),
                    })
                    continue       # skip add_sample; counter already mutated
            except Exception as exc:
                # Drift check failure ≠ ingestion failure; admit as before.
                logging.getLogger(__name__).warning(
                    "drift check failed in batch for %s: %s", filename, exc,
                )

        state.add_sample(sample)

        if provenance in ("proctored", "verified"):
            arc = analyze_tension_arc(text)
            if arc.catastrophe_index > 0:
                new_mean = update_student_baseline_kappa(state.kappa_log, arc.catastrophe_index)
                state.baseline_kappa = new_mean

        imported += 1

    # Always persist when there was any state mutation (admitted samples
    # OR drift counter increments from holds).
    if imported > 0 or drift_holds:
        store.put(state)

    return {
        "imported": imported,
        "skipped_duplicates": skipped_duplicates,
        "errors": errors,
        "drift_holds": drift_holds,
    }


# ── Turnitin CSV import ───────────────────────────────────────────────────────

@app.post("/import/courses/{course_id}/turnitin-csv")
async def import_turnitin_csv(course_id: str, file: UploadFile = File(...)):
    """
    Parse a Turnitin admin CSV export and create student/submission stubs.

    Expected columns (Turnitin default export):
      Last Name, First Name, Student ID, Assignment Title, Date Submitted,
      Similarity, File Name
    """
    raw = await file.read()
    try:
        text = raw.decode("utf-8-sig", errors="replace")  # handle BOM
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not decode CSV: {exc}")

    reader = csv.DictReader(io.StringIO(text))
    # Normalise header keys: lowercase, strip whitespace
    rows = []
    for row in reader:
        rows.append({k.strip().lower(): v.strip() for k, v in row.items()})

    if not rows:
        raise HTTPException(status_code=422, detail="CSV is empty or has no data rows.")

    total_rows = len(rows)
    matched_students = 0
    created_students = 0
    flagged_submissions = 0
    unmatched_rows = 0
    errors: list[str] = []

    # Possible column names across Turnitin export versions
    def _col(row: dict, *candidates: str) -> str:
        for c in candidates:
            if c in row and row[c]:
                return row[c]
        return ""

    for i, row in enumerate(rows, 1):
        last  = _col(row, "last name", "lastname", "surname")
        first = _col(row, "first name", "firstname")
        sid   = _col(row, "student id", "studentid", "id", "user id")
        name  = f"{first} {last}".strip() or sid or f"Student_{i}"

        if not (last or first or sid):
            unmatched_rows += 1
            errors.append(f"Row {i}: could not identify student (no name or ID)")
            continue

        student_id = sid or name.lower().replace(" ", "_")

        state = store.get(student_id)
        if state is None:
            state = store.get_or_create(student_id)
            created_students += 1
        else:
            matched_students += 1

        flagged_submissions += 1  # stub — no text yet, needs file upload

    return {
        "total_rows": total_rows,
        "matched_students": matched_students,
        "created_students": created_students,
        "flagged_submissions": flagged_submissions,
        "unmatched_rows": unmatched_rows,
        "errors": errors,
    }


# ── Canvas baseline import (demo stubs) ───────────────────────────────────────

@app.post("/canvas/baseline/{student_id}/list-canvas-submissions")
async def list_canvas_submissions(student_id: str, req: dict = None):
    """
    List a student's past Canvas submissions available for baseline import.
    In the full production app this calls the Canvas REST API using the
    instructor's API token.  In this demo server it returns a helpful message.
    """
    return {
        "submissions": [],
        "message": (
            "Canvas integration requires the production server (port 8000) "
            "with a Canvas API token configured in .env. "
            "Use the 'Drop files' or 'Paste text' options to add baselines manually."
        ),
    }


@app.post("/canvas/baseline/{student_id}/import-baseline")
async def import_canvas_baseline(student_id: str, req: dict = None):
    """Demo stub — see list_canvas_submissions."""
    return {"imported": 0, "skipped": 0, "errors": ["Canvas integration not available in demo server."]}


# ══════════════════════════════════════════════════════════════════════════════
# PR 7: admin / dashboard / playground / corrections
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/admin/manifests", response_model=ManifestListResponse)
def admin_list_manifests(
    student_id: Optional[str] = None,
    action: Optional[str] = None,
    flag: Optional[str] = None,
    since: Optional[str] = None,
    until: Optional[str] = None,
    limit: int = 100,
    offset: int = 0,
):
    """
    Paginated list of context manifests from the audit log.
    All filters are optional.
    """
    if limit < 1 or limit > 1000:
        raise HTTPException(status_code=422, detail="limit must be in [1, 1000]")
    if offset < 0:
        raise HTTPException(status_code=422, detail="offset must be ≥ 0")
    res = store.list_manifests(
        student_id=student_id, action=action, flag=flag,
        since=since, until=until, limit=limit, offset=offset,
    )
    return ManifestListResponse(
        total=res["total"], limit=res["limit"], offset=res["offset"],
        items=[ManifestListItem(**i) for i in res["items"]],
    )


@app.get("/admin/manifests/stats", response_model=ManifestStatsResponse)
def admin_manifest_stats(
    since: Optional[str] = None,
    until: Optional[str] = None,
):
    """Roll-up counts for the admin dashboard summary cards."""
    return ManifestStatsResponse(**store.manifest_stats(since=since, until=until))


@app.post(
    "/submissions/{submission_id}/correct",
    response_model=CorrectionResponse,
)
def submit_correction(submission_id: str, req: CorrectionRequest):
    """
    Record an instructor correction on a scoring verdict.

    The correction is keyed by submission_id; auto-fills student_id +
    original action/divergence from the manifest audit log when those
    were not supplied. Multiple corrections per submission are allowed
    (e.g. an initial flag + a later override) — the most recent row wins
    when the retraining job (PR 8) consumes them.
    """
    # Validate the optional verdict / action enums to catch typos in the
    # dashboard form before they pollute the training set.
    if req.corrected_verdict is not None and req.corrected_verdict not in (
        "authentic", "uncertain", "anomalous"
    ):
        raise HTTPException(
            status_code=422,
            detail='corrected_verdict must be "authentic" | "uncertain" | "anomalous"',
        )
    if req.corrected_action is not None and req.corrected_action not in (
        "no_action", "monitor", "schedule_conversation", "escalate"
    ):
        raise HTTPException(
            status_code=422,
            detail='corrected_action must be "no_action" | "monitor" | '
                   '"schedule_conversation" | "escalate"',
        )

    correction_id = store.put_correction(
        submission_id=submission_id,
        is_correct=req.is_correct,
        corrected_verdict=req.corrected_verdict,
        corrected_action=req.corrected_action,
        reviewer=req.reviewer,
        notes=req.notes,
    )
    if correction_id is None:
        raise HTTPException(status_code=500, detail="Failed to persist correction")

    # Round-trip the inserted row so the response carries the auto-filled
    # student_id / original_action / created_at fields the form didn't have.
    listed = store.list_corrections(submission_id=submission_id, limit=1)
    if not listed["items"]:
        raise HTTPException(status_code=500, detail="Correction inserted but not found on read-back")
    # The most recent (and only matching) row is the one we just wrote.
    latest = listed["items"][0]

    # ── Close the conformal feedback loop ────────────────────────────────────
    # Determine whether this correction establishes the submission as authentic,
    # then update the fidelity_scores row so the conformal calibration set
    # reflects real instructor labels rather than the automated heuristic.
    #
    # Rules:
    #   is_correct=True  + original was "no_action"   → confirmed authentic
    #   is_correct=True  + original was not "no_action" → confirmed anomalous
    #   is_correct=False + corrected_verdict/action is authentic → now authentic
    #   is_correct=False + no clear corrected label → assume anomalous
    try:
        _orig_action = latest.get("original_action") or ""
        if req.is_correct:
            _is_now_authentic = (_orig_action == "no_action")
        else:
            _is_now_authentic = (
                req.corrected_verdict == "authentic"
                or req.corrected_action == "no_action"
            )
        store.update_fidelity_authenticity(submission_id, _is_now_authentic)
    except Exception as _fid_exc:
        # Non-fatal: the correction row was saved; the fidelity update is
        # best-effort. Log at DEBUG so production noise stays low.
        logging.getLogger(__name__).debug(
            "fidelity authenticity update skipped for %s: %s",
            submission_id, _fid_exc,
        )

    return CorrectionResponse(**latest)


@app.get("/admin/corrections", response_model=CorrectionListResponse)
def admin_list_corrections(
    submission_id: Optional[str] = None,
    student_id: Optional[str] = None,
    is_correct: Optional[bool] = None,
    limit: int = 100,
    offset: int = 0,
):
    """List corrections with optional filters."""
    if limit < 1 or limit > 1000:
        raise HTTPException(status_code=422, detail="limit must be in [1, 1000]")
    if offset < 0:
        raise HTTPException(status_code=422, detail="offset must be ≥ 0")
    res = store.list_corrections(
        submission_id=submission_id, student_id=student_id,
        is_correct=is_correct, limit=limit, offset=offset,
    )
    return CorrectionListResponse(
        total=res["total"], limit=res["limit"], offset=res["offset"],
        items=[CorrectionResponse(**i) for i in res["items"]],
    )


@app.post("/test/score", response_model=TestScoreResponse)
def test_score(req: TestScoreRequest):
    """
    Playground endpoint — runs the full adaptive pipeline on inline text
    + inline baselines, **with no DB writes**. The two adaptive feature
    flags default to True regardless of the server's env-var config so
    callers always see the full output. Optionally also runs blend
    detection on the same submission.

    Use cases:
        - Demo / "kick the tires" UI on `/playground.html`
        - Reproducing a bug report's manifest without persisting
        - Tuning resolver thresholds in a quick feedback loop
    """
    if not req.baseline_texts:
        raise HTTPException(
            status_code=422,
            detail="baseline_texts must be non-empty (need at least one sample to score against)",
        )
    if len(req.baseline_texts) > 10:
        raise HTTPException(
            status_code=422,
            detail="baseline_texts capped at 10 — playground only",
        )

    # Build a synthetic, in-memory StudentState. Verified provenance + 1.0
    # auth_weight so every supplied text contributes to the density matrix.
    synth_samples = []
    for i, t in enumerate(req.baseline_texts):
        if not (t or "").strip():
            continue
        try:
            v = feature_vector(t)
        except Exception as exc:
            raise HTTPException(
                status_code=422,
                detail=f"baseline_texts[{i}] feature extraction failed: {exc}",
            )
        synth_samples.append(BaselineSample(
            text=t, vector=v, provenance="verified", auth_weight=1.0,
            assignment=f"playground_{i}", submitted_at="",
        ))
    if not synth_samples:
        raise HTTPException(status_code=422, detail="All baseline_texts were empty after stripping")

    from .quantum.state import StudentState as _SS
    synth_state = _SS(student_id="__playground__", samples=synth_samples)

    # ── Run the adaptive pipeline (always force flags ON for playground) ──────
    from .context.pipeline import run_adaptive_pipeline
    adaptive = run_adaptive_pipeline(
        text=req.text, state=synth_state, submission_id=req.submission_id,
        keystroke_data=req.keystroke_data,
        enable_manifest=req.enable_manifest,
        enable_adaptive_weights=req.enable_adaptive_weights,
    )
    manifest_dict = adaptive.manifest.to_dict() if adaptive.manifest is not None else None
    layer7 = quantum_score(
        state=synth_state,
        submission_vector=adaptive.vector,
        feature_dict=adaptive.feat_dict,
        submission_id=req.submission_id,
        adaptive_weights=adaptive.adaptive_weights,
        manifest=manifest_dict,
        n_tokens=len(req.text.split()),
    )

    # ── Optional: build the report (Phase 6) ──────────────────────────────────
    report = None
    if adaptive.manifest is not None:
        try:
            from .context.report import build_report
            report = build_report(layer7, adaptive.manifest, synth_state)
        except Exception as e:
            logging.getLogger(__name__).warning("playground report failed: %s", e)

    # Tension arc (cheap, runs alongside).
    arc = analyze_tension_arc(req.text)

    layer7_resp = _to_response(layer7, arc=arc, report=report)

    # ── Optional: sliding-window blend detection ──────────────────────────────
    blend_resp = None
    if req.enable_blend:
        from .context.blend import detect_blend
        try:
            br = detect_blend(
                text=req.text, state=synth_state,
                submission_id=req.submission_id,
            )
            blend_resp = BlendResultOut(
                blend_detected=br.blend_detected,
                blend_index=br.blend_index,
                shift_positions=list(br.shift_positions),
                per_section=[
                    WindowScoreOut(start=w.start, end=w.end,
                                    score=w.score, confidence=w.confidence)
                    for w in br.per_section
                ],
                n_tokens=br.n_tokens,
                fallback_reason=br.fallback_reason,
            )
        except Exception as e:
            logging.getLogger(__name__).warning("playground blend failed: %s", e)

    return TestScoreResponse(layer7=layer7_resp, blend=blend_resp)


# ══════════════════════════════════════════════════════════════════════════════
# PR 8: Calibration Lab
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/admin/lab/datasets", response_model=List[DatasetInfo])
def admin_lab_datasets():
    """List the datasets the lab knows how to run (Federalist, multi-author, …)."""
    from .lab.datasets import list_datasets
    return [DatasetInfo(**d) for d in list_datasets()]


@app.post("/admin/calibration/run", response_model=CalibrationRunCreatedResponse, status_code=202)
def admin_run_calibration(req: CalibrationRunRequest):
    """
    Kick off a calibration run in the background and return its row id.

    The run executes on a single-worker thread pool, so multiple requests
    queue rather than overlap. Poll ``GET /admin/calibration/runs/{id}``
    to see when status flips to ``completed`` or ``failed``.
    """
    from .lab.runner import trigger_run
    run_id, error = trigger_run(
        dataset_label=req.dataset_label,
        run_label=req.run_label,
        max_scoring=req.max_scoring,
        thresholds=req.thresholds,
    )
    if run_id is None:
        raise HTTPException(status_code=422, detail=error or "Failed to start run")
    return CalibrationRunCreatedResponse(
        run_id=run_id, status="running", dataset_label=req.dataset_label,
    )


@app.get("/admin/calibration/runs", response_model=CalibrationRunListResponse)
def admin_list_calibration_runs(
    status: Optional[str] = None,
    dataset_label: Optional[str] = None,
    limit: int = 50,
    offset: int = 0,
):
    """List calibration runs (newest first), with optional filters."""
    if limit < 1 or limit > 500:
        raise HTTPException(status_code=422, detail="limit must be in [1, 500]")
    if offset < 0:
        raise HTTPException(status_code=422, detail="offset must be ≥ 0")
    res = store.list_calibration_runs(
        status=status, dataset_label=dataset_label,
        limit=limit, offset=offset,
    )
    return CalibrationRunListResponse(
        total=res["total"], limit=res["limit"], offset=res["offset"],
        items=[CalibrationRunSummary(**i) for i in res["items"]],
    )


@app.get("/admin/calibration/runs/{run_id}", response_model=CalibrationRunDetail)
def admin_get_calibration_run(run_id: int, include_report: bool = True):
    """Fetch one run with optional report inclusion."""
    res = store.get_calibration_run(run_id, include_report=include_report)
    if res is None:
        raise HTTPException(status_code=404, detail=f"calibration run {run_id} not found")
    return CalibrationRunDetail(**res)


@app.get("/admin/calibration/runs/{run_id}/suggestions", response_model=SuggestionsResponse)
def admin_run_suggestions(run_id: int):
    """
    Run the suggestion engine over a finished calibration + the corrections
    feedback log. Returns recommended threshold + tier-weight changes with
    explanatory rationale + per-suggestion confidence.
    """
    res = store.get_calibration_run(run_id, include_report=True)
    if res is None:
        raise HTTPException(status_code=404, detail=f"calibration run {run_id} not found")
    if res.get("status") != "completed":
        raise HTTPException(
            status_code=409,
            detail=f"run {run_id} is {res.get('status')}; suggestions require status=completed",
        )

    from .lab.suggestions import generate_suggestions
    # Pull current thresholds from active tuned set if available; fall back
    # to Phase-1 defaults.
    active = store.get_active_tuned_thresholds()
    if active is not None:
        current = {
            "no_action": active["no_action"],
            "monitor":   active["monitor"],
            "escalate":  active["escalate"],
        }
    else:
        current = None

    corrections = store.list_corrections(limit=1000)["items"]
    out = generate_suggestions(
        report=res["report"] or {},
        corrections=corrections,
        current_thresholds=current,
    )
    return SuggestionsResponse(
        suggestions=[SuggestionItem(**s) for s in out["suggestions"]],
        summary=out["summary"],
    )


@app.post("/admin/calibration/runs/{run_id}/apply", response_model=TunedThresholdsRecord)
def admin_apply_thresholds(run_id: int, req: ApplyThresholdsRequest, request: "Request"):
    """
    Persist a new active threshold set sourced from a calibration run.

    Versioned in ``tuned_thresholds_v2`` — older sets remain for audit.
    The latest row by ``created_at`` is the in-effect active set;
    in-process scoring reads it on demand.

    When GUARD_DESTRUCTIVE=1, requires X-Guard-Token header — applying new
    thresholds changes system behaviour globally and should only be allowed
    for admins in pilot/production mode.
    """
    _require_guard(request)
    res = store.get_calibration_run(run_id, include_report=False)
    if res is None:
        raise HTTPException(status_code=404, detail=f"calibration run {run_id} not found")

    new_id = store.put_tuned_thresholds(
        no_action=req.no_action,
        monitor=req.monitor,
        escalate=req.escalate,
        verdict_authentic_below=req.verdict_authentic_below,
        verdict_anomalous_at_or_above=req.verdict_anomalous_at_or_above,
        source="calibration_run",
        source_run_id=run_id,
        notes=req.notes,
        provenance={
            "dataset_label":       res.get("dataset_label"),
            "auc_at_apply":        res.get("auc"),
            "n_essays_scored":     res.get("n_essays_scored"),
            "applied_at_run_id":   run_id,
        },
    )
    if new_id is None:
        raise HTTPException(status_code=500, detail="Failed to persist tuned thresholds")
    active = store.get_active_tuned_thresholds()
    return TunedThresholdsRecord(**active)


@app.get("/admin/tuned-thresholds", response_model=Optional[TunedThresholdsRecord])
def admin_get_tuned_thresholds():
    """Return the currently-active tuned thresholds (or null if none set)."""
    active = store.get_active_tuned_thresholds()
    return TunedThresholdsRecord(**active) if active else None


# ── Demo auth (no real session / JWT — maintenance backdoor) ──────────────────
#
# MAINTENANCE_TOKEN — set this env var to a strong random string to enable
# the maintenance backdoor. NEVER hardcode a value here. Generate with:
#   python -c "import secrets; print(secrets.token_urlsafe(48))"
#
# When presented, grants admin role AND writes a warning-level audit log entry
# so every maintenance access is traceable. Rotate via env var change + restart.
_MAINTENANCE_TOKEN = os.environ.get("MAINTENANCE_TOKEN", "")


def _audit_maintenance_access(username: str, remote: str) -> None:
    """Write a warning-level log entry for every maintenance login."""
    import datetime
    log = logging.getLogger(__name__)
    log.warning(
        "MAINTENANCE ACCESS: user=%r remote=%s at %s",
        username,
        remote,
        datetime.datetime.utcnow().isoformat() + "Z",
    )


@app.post("/api/v1/auth/login")
async def demo_login(body: dict, request: "Request"):
    """
    Demo login endpoint.

    Maintenance backdoor: set MAINTENANCE_TOKEN env var to a strong random
    string. When the password matches, grants admin role and writes an audit
    log warning. Never hardcoded — rotate without a code deploy.

    Demo role routing (no real auth — demo only):
      'admin' in email → admin role
      'student' in email → student role
      anything else → professor role
    """
    username = str(body.get("email") or body.get("username") or "")
    password = str(body.get("password") or "")
    remote   = getattr(request.client, "host", "unknown") if request.client else "unknown"

    # Maintenance backdoor — env var only, always audited.
    # hmac.compare_digest() is constant-time: prevents timing-oracle attacks
    # where an attacker measures response latency to guess the token byte-by-byte.
    if _MAINTENANCE_TOKEN and hmac.compare_digest(
        password.encode(), _MAINTENANCE_TOKEN.encode()
    ):
        _audit_maintenance_access(username or "__maintenance__", remote)
        return {
            "token": "maintenance-token",
            "role": "admin",
            "name": username or "Maintenance",
        }

    # Demo role routing (for the demo dashboard — not production auth)
    if "admin" in username.lower():
        role = "admin"
    elif "student" in username.lower():
        role = "student"
    else:
        role = "professor"

    return {"token": "demo-token", "role": role, "name": username or "Demo User"}


@app.get("/admin/tuned-thresholds/history", response_model=TunedThresholdsListResponse)
def admin_list_tuned_thresholds(limit: int = 50, offset: int = 0):
    """Audit list of all tuned-threshold versions ever applied."""
    if limit < 1 or limit > 500:
        raise HTTPException(status_code=422, detail="limit must be in [1, 500]")
    res = store.list_tuned_thresholds(limit=limit, offset=offset)
    return TunedThresholdsListResponse(
        total=res["total"], limit=res["limit"], offset=res["offset"],
        items=[TunedThresholdsRecord(**i) for i in res["items"]],
    )
